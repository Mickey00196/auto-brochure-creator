"""§20 Export — CSV / Excel / JSON / Word renderings of the comparison table.
PDF/PPTX/one-pager live in app/services/brochure/.
"""
from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

from app.models import Proposal
from app.services.comparison import build_comparison_table

_HEADERS = [
    "Location", "Floor", "Area (m2)", "Rent EUR/m2/yr", "Service Charge EUR/m2/yr",
    "All-in EUR/m2/yr", "Estimated Annual Cost EUR", "Contract Term", "Address",
]


def _rows_as_lists(proposal: Proposal) -> list[list]:
    rows = build_comparison_table(proposal.selected_units)
    out = []
    for r in rows:
        out.append([
            r.building_name,
            r.floor or "",
            r.available_area_m2,
            r.rent_eur_per_m2_year if r.rent_eur_per_m2_year is not None else "TBD",
            r.service_charge_eur_per_m2_year if r.service_charge_eur_per_m2_year is not None else "TBD",
            r.all_in_rate_eur_per_m2_year if r.all_in_rate_eur_per_m2_year is not None else "TBD",
            r.estimated_annual_cost_eur if r.estimated_annual_cost_eur is not None else "TBD",
            r.contract_term or "TBD",
            r.address,
        ])
    return out


def export_csv(proposal: Proposal, out_path: Path) -> Path:
    with out_path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(_HEADERS)
        writer.writerows(_rows_as_lists(proposal))
    return out_path


def export_json(proposal: Proposal, out_path: Path) -> Path:
    rows = build_comparison_table(proposal.selected_units)
    payload = {
        "proposal_id": proposal.proposal_id,
        "title": proposal.title,
        "client": proposal.client.company_name,
        "units": [r.to_dict() for r in rows],
    }
    out_path.write_text(json.dumps(payload, indent=2, default=str), encoding="utf-8")
    return out_path


def export_excel(proposal: Proposal, out_path: Path) -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "Comparison"
    ws.append(_HEADERS)
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="0B0B0C", end_color="0B0B0C", fill_type="solid")
    for row in _rows_as_lists(proposal):
        ws.append(row)
    for column_cells in ws.columns:
        length = max(len(str(c.value)) for c in column_cells)
        ws.column_dimensions[column_cells[0].column_letter].width = min(max(length + 2, 10), 40)
    wb.save(out_path)
    return out_path


def export_word(proposal: Proposal, out_path: Path) -> Path:
    doc = Document()
    title = doc.add_heading(proposal.title, level=0)
    doc.add_paragraph(f"Prepared for {proposal.client.company_name} by {proposal.prepared_by or 'your advisor'}")

    rows = build_comparison_table(proposal.selected_units)
    table = doc.add_table(rows=1, cols=len(_HEADERS))
    table.style = "Light Grid Accent 1"
    for i, header in enumerate(_HEADERS):
        table.rows[0].cells[i].text = header

    for values in _rows_as_lists(proposal):
        cells = table.add_row().cells
        for i, val in enumerate(values):
            cells[i].text = str(val)

    for style_name in ("Normal",):
        doc.styles[style_name].font.size = Pt(10)

    doc.save(out_path)
    return out_path
